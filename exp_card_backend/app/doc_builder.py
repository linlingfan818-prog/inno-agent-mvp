from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt

from .schemas import ExperimentCard
from typing import List, Dict


def _add_bullets(document: Document, items: List[str]) -> None:
    if not items:
        p = document.add_paragraph()
        p.add_run("无").italic = True
        return
    for item in items:
        document.add_paragraph(item, style="List Bullet")



def build_experiment_card_docx(card: ExperimentCard, file_path: Path) -> None:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    document.add_heading("创新实验卡", level=0)
    document.add_paragraph(f"项目名称：{card.project_name}")

    document.add_heading("1. 核心假设", level=1)
    document.add_paragraph(card.core_hypothesis)

    document.add_heading("2. OKRs / KRs 与核心假设映射", level=1)
    if card.hypothesis_mapping:
        table = document.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "OKR / KR"
        hdr[1].text = "核心假设"
        hdr[2].text = "是否可行"
        for row in card.hypothesis_mapping:
            cells = table.add_row().cells
            cells[0].text = row.get("okr_or_kr", "")
            cells[1].text = row.get("hypothesis", "")
            cells[2].text = row.get("feasibility_check", "")
    else:
        document.add_paragraph("无")

    document.add_heading("3. 实验周期", level=1)
    document.add_paragraph(card.experiment_cycle)

    document.add_heading("4. 实验方法", level=1)
    document.add_paragraph(card.experiment_method)

    document.add_heading("5. 目标用户", level=1)
    _add_bullets(document, card.target_users)

    document.add_heading("6. WHY / WHAT / VALUE", level=1)
    document.add_paragraph(f"WHY: {card.why_statement}")
    document.add_paragraph(f"WHAT: {card.what_solution}")
    document.add_paragraph(f"VALUE: {card.value_statement}")

    document.add_heading("7. 实验步骤", level=1)
    _add_bullets(document, card.experiment_steps)

    document.add_heading("8. 成功指标", level=1)
    _add_bullets(document, card.success_metrics)

    document.add_heading("9. 风险与注意事项", level=1)
    _add_bullets(document, card.risks_and_watchouts)

    document.add_heading("10. 项目完成状态评估清单", level=1)
    _add_bullets(document, card.completion_checklist)

    if card.critical_acceptance_standard:
        document.add_heading("11. 关键验收标准 (Critical Acceptance Standard)", level=1)
        
        document.add_heading("验收环境与前提", level=2)
        document.add_paragraph(card.critical_acceptance_standard.environment_and_prerequisites)
        
        document.add_heading("核心通关指标 (Must-have)", level=2)
        _add_bullets(document, card.critical_acceptance_standard.must_have_metrics)
        
        document.add_heading("一票否决项 (Red lines)", level=2)
        _add_bullets(document, card.critical_acceptance_standard.red_lines)

    document.add_heading("12. 总结", level=1)
    document.add_paragraph(card.output_summary)

    document.save(str(file_path))
